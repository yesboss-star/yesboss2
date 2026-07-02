from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Inter"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(4, 44, 83)

# ── Title page ──────────────────────────────────────────
doc.add_heading("YesBoss — Google & Outlook Integration Plan", level=0)
p = doc.add_paragraph(f"Prepared: {datetime.date.today().isoformat()}")
p.runs[0].font.color.rgb = RGBColor(107, 106, 99)
doc.add_paragraph("")

# ── 1. Overview ─────────────────────────────────────────
doc.add_heading("1. Overview", level=1)
doc.add_paragraph(
    "This document outlines the work required to add Google Calendar + Gmail and "
    "Outlook (Microsoft 365) Calendar + Mail integrations to YesBoss, following "
    "the same OAuth 2.0 pattern already implemented for Zoho."
)
doc.add_paragraph(
    "Both integrations allow users to connect their personal/work accounts via "
    "OAuth 2.0, after which they can read calendars, create events, and optionally "
    "interact with mail — all from within the YesBoss dashboard."
)

# ── 2. Architecture ─────────────────────────────────────
doc.add_heading("2. Architecture (Same Pattern as Zoho)", level=1)
doc.add_paragraph(
    "The Zoho integration follows a clean OAuth 2.0 authorization code flow with "
    "PKCE-equivalent state parameter. Both Google and Outlook will mirror this "
    "exactly."
)

doc.add_heading("Frontend (Zustand Store + React Component)", level=2)
items = [
    ("Zustand store", "Manages state: connected, email, loading, connecting. Exposes checkStatus(), connect(), disconnect()."),
    ("Connect button", "Calls GET /auth-url, opens OAuth popup, polls for popup close, calls checkStatus()."),
    ("Calendar booking UI", "Existing ZohoCalendarBooking component refactored to accept any provider."),
]
for name, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Backend (FastAPI + Python)", level=2)
items = [
    ("OAuth handler class", "Generates auth URL, exchanges code for tokens, refresh flow, token CRUD in MongoDB."),
    ("API routes", "/auth-url, /callback, /status, /disconnect per provider."),
    ("Calendar API class", "List calendars, list events, check free/busy, create events via provider's REST API."),
    ("Token storage", "Each provider gets its own MongoDB collection: google_tokens, outlook_tokens."),
]
for name, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

# ── 3. Google Integration ───────────────────────────────
doc.add_heading("3. Google Integration", level=1)

doc.add_heading("3.1 Prerequisites", level=2)
doc.add_paragraph("From the Google Cloud Console (https://console.cloud.google.com):")

table = doc.add_table(rows=5, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Item", "Where", "Cost"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ("Google Cloud Project", "console.cloud.google.com → Create Project", "Free"),
    ("OAuth 2.0 Web Client ID", "APIs & Services → Credentials → Create Credentials", "Free"),
    ("Google Calendar API enabled", "APIs & Services → Library → Search 'Google Calendar API'", "Free (1M req/day)"),
    ("Gmail API enabled (optional)", "APIs & Services → Library → Search 'Gmail API'", "Free"),
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph("")
doc.add_heading("3.2 Required OAuth Scopes", level=2)
scopes = [
    "https://www.googleapis.com/auth/calendar — Full calendar read/write",
    "https://www.googleapis.com/auth/calendar.events — Read/create events (narrower)",
    "https://www.googleapis.com/auth/gmail.readonly — Read emails (optional)",
    "https://www.googleapis.com/auth/gmail.send — Send emails (optional)",
]
for s in scopes:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3.3 Endpoints Used", level=2)
endpoints = [
    ("https://accounts.google.com/o/oauth2/v2/auth", "Authorization URL"),
    ("https://oauth2.googleapis.com/token", "Token exchange & refresh"),
    ("https://www.googleapis.com/calendar/v3/calendars/primary", "Calendar API"),
    ("https://gmail.googleapis.com/gmail/v1/users/me/messages", "Gmail API (optional)"),
]
for url, desc in endpoints:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{url}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(15, 110, 86)
    p.add_run(f" — {desc}")

doc.add_heading("3.4 Files to Create", level=2)
table = doc.add_table(rows=7, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Layer", "File", "Purpose"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
files = [
    ("Backend", "app/core/google/__init__.py", "Package init"),
    ("Backend", "app/core/google/base.py", "GoogleOAuth class (auth URL, exchange, refresh, token CRUD)"),
    ("Backend", "app/core/google/calendar.py", "GoogleCalendar class (list, get, create events)"),
    ("Backend", "app/api/google_auth.py", "Routes: /google/auth-url, /google/callback, /google/status, /google/disconnect"),
    ("Frontend", "stores/googleStore.ts", "Zustand store (checkStatus, connect, disconnect)"),
    ("Frontend", "components/owners/GoogleConnectButton.tsx", "Connect/disconnect button UI"),
]
for r, (layer, filepath, purpose) in enumerate(files):
    table.rows[r + 1].cells[0].text = layer
    table.rows[r + 1].cells[1].text = filepath
    table.rows[r + 1].cells[2].text = purpose

doc.add_paragraph("")
doc.add_heading("3.5 Environment Variables", level=2)
env_vars = [
    ("GOOGLE_CLIENT_ID", "xxx.apps.googleusercontent.com"),
    ("GOOGLE_CLIENT_SECRET", "GOCSPX-xxxxxxxxxxxx"),
    ("GOOGLE_REDIRECT_URI", "http://localhost:8000/api/v1/google/callback"),
]
for var, example in env_vars:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{var}")
    run.bold = True
    p.add_run(f" — Example: {example}")

# ── 4. Outlook Integration ──────────────────────────────
doc.add_heading("4. Outlook (Microsoft 365) Integration", level=1)

doc.add_heading("4.1 Prerequisites", level=2)
doc.add_paragraph("From the Microsoft Azure portal (https://portal.azure.com):")

table = doc.add_table(rows=5, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Item", "Where", "Cost"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ("Azure subscription", "portal.azure.com", "Free (no spend required)"),
    ("App registration", "Microsoft Entra ID → App registrations → New registration", "Free"),
    ("Client secret", "Certificates & Secrets → New client secret", "Free"),
    ("M365 dev sandbox (testing)", "developer.microsoft.com/microsoft-365/dev-program", "Free (90-day renewable)"),
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r + 1].cells[c].text = val

doc.add_paragraph("")
doc.add_heading("4.2 App Registration Configuration", level=2)
config_items = [
    "Redirect URI type: Web, URI: http://localhost:8000/api/v1/outlook/callback",
    "Supported account types: 'Accounts in any organizational directory (Multitenant)' + 'Personal Microsoft accounts'",
    "API Permissions: Add Microsoft Graph → Delegated permissions",
]
for item in config_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.3 Required Microsoft Graph Scopes", level=2)
scopes = [
    "Calendars.ReadWrite — Read and create calendar events",
    "User.Read — Get display name and email",
    "offline_access — Receive refresh tokens (mandatory for long-lived access)",
    "Mail.ReadWrite — Read and send emails (optional)",
]
for s in scopes:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("4.4 Endpoints Used", level=2)
endpoints = [
    ("https://login.microsoftonline.com/common/oauth2/v2.0/authorize", "Authorization URL"),
    ("https://login.microsoftonline.com/common/oauth2/v2.0/token", "Token exchange & refresh"),
    ("https://graph.microsoft.com/v1.0/me/calendar/events", "Calendar events (Microsoft Graph)"),
    ("https://graph.microsoft.com/v1.0/me/messages", "Mail messages (Graph, optional)"),
]
for url, desc in endpoints:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{url}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(15, 110, 56)
    p.add_run(f" — {desc}")

doc.add_heading("4.5 Files to Create", level=2)
table = doc.add_table(rows=7, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Layer", "File", "Purpose"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
files = [
    ("Backend", "app/core/outlook/__init__.py", "Package init"),
    ("Backend", "app/core/outlook/base.py", "OutlookOAuth class (same pattern as ZohoOAuth)"),
    ("Backend", "app/core/outlook/calendar.py", "OutlookCalendar via Microsoft Graph API"),
    ("Backend", "app/api/outlook_auth.py", "Routes: /outlook/auth-url, /outlook/callback, /outlook/status, /outlook/disconnect"),
    ("Frontend", "stores/outlookStore.ts", "Zustand store"),
    ("Frontend", "components/owners/OutlookConnectButton.tsx", "Connect/disconnect button UI"),
]
for r, (layer, filepath, purpose) in enumerate(files):
    table.rows[r + 1].cells[0].text = layer
    table.rows[r + 1].cells[1].text = filepath
    table.rows[r + 1].cells[2].text = purpose

doc.add_paragraph("")
doc.add_heading("4.6 Environment Variables", level=2)
env_vars = [
    ("OUTLOOK_CLIENT_ID", "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx (Application ID from Azure)"),
    ("OUTLOOK_CLIENT_SECRET", "Client secret value from Certificates & Secrets"),
    ("OUTLOOK_TENANT_ID", "common (for multi-tenant) or your specific tenant ID"),
    ("OUTLOOK_REDIRECT_URI", "http://localhost:8000/api/v1/outlook/callback"),
]
for var, example in env_vars:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{var}")
    run.bold = True
    p.add_run(f" — {example}")

# ── 5. Implementation Steps ─────────────────────────────
doc.add_heading("5. Implementation Steps", level=1)

doc.add_heading("Phase 1: Google Calendar (estimated: 4-6 hours)", level=2)
phase1 = [
    "Create app/core/google/ package with OAuth base class",
    "Create app/api/google_auth.py with /google/* routes",
    "Create app/core/google/calendar.py with Calendar API methods",
    "Create frontend googleStore.ts and GoogleConnectButton.tsx",
    "Register in main.py router and frontend settings page",
    "Test: connect, list events, create event, disconnect",
]
for s in phase1:
    doc.add_paragraph(s, style="List Number")

doc.add_heading("Phase 2: Outlook Calendar (estimated: 6-8 hours)", level=2)
phase2 = [
    "Register app in Azure portal, configure permissions",
    "Create app/core/outlook/ package with OAuth base class",
    "Create app/api/outlook_auth.py with /outlook/* routes",
    "Create app/core/outlook/calendar.py with Graph API methods",
    "Create frontend outlookStore.ts and OutlookConnectButton.tsx",
    "Register in main.py router and frontend settings page",
    "Test: connect, list events, create event, disconnect",
]
for s in phase2:
    doc.add_paragraph(s, style="List Number")

doc.add_heading("Phase 3 (Optional): Refactor to Unified Provider Model", level=2)
phase3 = [
    "Create a common OAuthBase class that Zoho, Google, and Outlook inherit from",
    "Create a unified CalendarProvider interface so booking UI works with any provider",
    "Let users connect multiple providers simultaneously and select which to use",
    "Add a provider selector dropdown to ZohoCalendarBooking (now CalendarBooking)",
]
for s in phase3:
    doc.add_paragraph(s, style="List Number")

# ── 6. Files to Modify ──────────────────────────────────
doc.add_heading("6. Existing Files to Modify", level=1)
mods = [
    ("backend/app/main.py", "Add google_auth_router and outlook_auth_router includes"),
    ("backend/app/core/config.py", "Add Google and Outlook env vars"),
    ("backend/.env.example", "Add Google and Outlook variable stubs"),
    ("frontend/src/app/dashboard/settings/page.tsx", "Add GoogleConnectButton and OutlookConnectButton"),
]
table = doc.add_table(rows=5, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "File"
table.rows[0].cells[1].text = "Change"
for r, (filepath, change) in enumerate(mods):
    table.rows[r + 1].cells[0].text = filepath
    table.rows[r + 1].cells[1].text = change

# ── 7. Key Differences Reference ────────────────────────
doc.add_heading("7. OAuth Provider Comparison", level=1)
table = doc.add_table(rows=7, cols=4)
table.style = "Light Grid Accent 1"
headers = ["Detail", "Zoho", "Google", "Outlook"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ("Auth URL", "accounts.zoho.com/oauth/v2/auth", "accounts.google.com/o/oauth2/v2/auth", "login.microsoftonline.com/common/oauth2/v2.0/authorize"),
    ("Token URL", "accounts.zoho.com/oauth/v2/token", "oauth2.googleapis.com/token", "login.microsoftonline.com/common/oauth2/v2.0/token"),
    ("Auth header", "Zoho-oauthtoken {token}", "Bearer {token}", "Bearer {token}"),
    ("Calendar API", "calendar.zoho.com/api/v1", "www.googleapis.com/calendar/v3", "graph.microsoft.com/v1.0/me/calendar"),
    ("Refresh token", "Automatic", "Automatic", "Requires offline_access scope"),
    ("MongoDB collection", "zoho_tokens", "google_tokens", "outlook_tokens"),
]
for r, (detail, zoho, google, outlook) in enumerate(data):
    table.rows[r + 1].cells[0].text = detail
    table.rows[r + 1].cells[1].text = zoho
    table.rows[r + 1].cells[2].text = google
    table.rows[r + 1].cells[3].text = outlook

# ── Save ────────────────────────────────────────────────
doc.save("C:\\VSLLP\\krisha\\2\\yesboss2\\yesboss_calendar_integration_plan.docx")
print("Document saved: yesboss_calendar_integration_plan.docx")
