from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Inter"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(4, 44, 83)

doc.add_heading("YesBoss — Document Requirements", level=0)
p = doc.add_paragraph(f"Prepared: {datetime.date.today().isoformat()}")
p.runs[0].font.color.rgb = RGBColor(107, 106, 99)
doc.add_paragraph("")

doc.add_heading("1. Documents Required to Onboard a Company (Owner)", level=1)
doc.add_paragraph(
    "These documents are collected during the owner/company onboarding flow. They help YesBoss "
    "understand the business, its industry, structure, products, operations, and strategy. "
    "The system also auto-detects industry from the company domain and suggests relevant documents."
)

doc.add_heading("1.1 Generic Documents (All Industries)", level=2)
table = doc.add_table(rows=10, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["#", "Document", "Purpose"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ("1", "Company Overview / About Us", "Company identity, vision, mission, founding story, core values"),
    ("2", "Organization Chart & Team List", "Hierarchy, team structure, reporting lines, department mapping"),
    ("3", "List of Products or Services", "Core offerings, product lines, service catalog, revenue drivers"),
    ("4", "Standard Operating Procedures (SOPs)", "Process documentation, operational workflows, compliance steps"),
    ("5", "Pricing Sheet or Rate Card", "Pricing strategy, billing rates, cost structure, margins"),
    ("6", "Customer or Vendor Contracts", "Key relationships, terms, obligations, revenue commitments"),
    ("7", "Marketing or Sales Playbook", "Go-to-market strategy, sales process, lead generation, channels"),
    ("8", "Goals, OKRs, or Quarterly Plans", "Strategic objectives, key results, timelines, ownership"),
    ("9", "Company Notes / Free Text", "Any additional context the owner wants to share"),
]
for r, (num, doc_name, purpose) in enumerate(data):
    table.rows[r + 1].cells[0].text = num
    table.rows[r + 1].cells[1].text = doc_name
    table.rows[r + 1].cells[2].text = purpose

doc.add_paragraph("")
doc.add_heading("1.2 Industry-Specific Documents", level=2)
doc.add_paragraph(
    "Based on the industry detected from the company domain or selected during onboarding, "
    "YesBoss suggests additional relevant documents. Below are key examples:"
)

industry_data = [
    ("Technology & Software", "Technical architecture docs, Product roadmaps, Tech stack documentation, Software requirements specifications, API documentation"),
    ("Banking & Financial Services", "Financial statements (P&L, Balance Sheet), Cash flow reports, Risk assessment documents, Customer portfolio data"),
    ("Manufacturing & Industrial", "Production schedules, Quality control reports, Equipment maintenance logs, Safety compliance documentation"),
    ("Healthcare & Life Sciences", "Patient flow diagrams, Staff scheduling templates, Medical inventory records, Compliance documentation"),
    ("Legal Services & Law", "Case files, Legal research documents, Client agreements, Compliance checklists"),
    ("Accounting & Taxation", "Financial statements, Tax filing documents, Audit reports, Compliance checklists"),
    ("E-commerce & Online Retail", "Product catalogs, Sales analytics reports, Customer segmentation data, Inventory management sheets"),
    ("Real Estate & Property", "Property listings, Lease agreements, Maintenance schedules, Market analysis reports"),
    ("Hospitality & Travel", "Booking records, Revenue management reports, Staff rosters, Guest satisfaction surveys"),
    ("Education & Training", "Curriculum documents, Student enrollment data, Staff credentials, Accreditation records"),
]

table = doc.add_table(rows=len(industry_data) + 1, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Industry", "Suggested Documents", "Value"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, (industry, docs) in enumerate(industry_data):
    table.rows[r + 1].cells[0].text = industry
    table.rows[r + 1].cells[1].text = docs
    table.rows[r + 1].cells[2].text = "AI uses these to generate accurate default goals, KPIs, and dashboards"

doc.add_paragraph("")
doc.add_heading("1.3 Auto-Detected Information", level=2)
auto_items = [
    "Company name and domain — extracted from signup email domain",
    "Industry and micro-vertical — AI-suggested from domain analysis",
    "Social media links — detected from web presence",
    "Company size — estimated from domain data and employee count",
    "Owner persona — captured via AI questionnaire during onboarding",
]
for item in auto_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("")
doc.add_heading("2. Documents Required for Final Results (Org Health & Growth)", level=1)
doc.add_paragraph(
    "These documents are consumed and generated by YesBoss to produce organizational health scores, "
    "financial insights, performance reports, and growth recommendations. They form the foundation "
    "of the AI Business Operating System's analytics engine."
)

doc.add_heading("2.1 Financial Documents (P&L, Cash Flow, Financial Health)", level=2)
table = doc.add_table(rows=8, cols=3)
table.style = "Light Grid Accent 1"
headers = ["#", "Document", "How It Is Used"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
fin_data = [
    ("1", "Profit & Loss Statement (P&L)", "Revenue, expenses, profitability analysis, cost structure"),
    ("2", "Balance Sheet", "Assets, liabilities, equity, financial stability assessment"),
    ("3", "Cash Flow Reports", "Liquidity analysis, burn rate, runway calculation"),
    ("4", "Revenue Breakdown / Run Rate", "ARR/MRR calculation, growth rate, revenue forecasting"),
    ("5", "Tax Filing Documents", "Compliance status, audit readiness assessment"),
    ("6", "Audit Reports", "Financial integrity, risk assessment, internal controls"),
    ("7", "Pricing Sheet / Rate Card", "Margin analysis, pricing strategy optimization"),
]
for r, (num, doc_name, usage) in enumerate(fin_data):
    table.rows[r + 1].cells[0].text = num
    table.rows[r + 1].cells[1].text = doc_name
    table.rows[r + 1].cells[2].text = usage

doc.add_paragraph("")
doc.add_heading("2.2 Operational Documents (SOPs & Efficiency)", level=2)
table = doc.add_table(rows=7, cols=3)
table.style = "Light Grid Accent 1"
headers = ["#", "Document", "How It Is Used"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
ops_data = [
    ("1", "Standard Operating Procedures (SOPs)", "Workflow pattern recording, efficiency scoring, bottleneck detection"),
    ("2", "Workflow Documentation", "Continuous learning engine input, automation opportunity identification"),
    ("3", "Production Schedules", "Capacity planning, utilization analysis, throughput optimization"),
    ("4", "Quality Control Reports", "Defect rate analysis, quality score, process improvement suggestions"),
    ("5", "Equipment Maintenance Logs", "Downtime analysis, preventive maintenance scheduling"),
    ("6", "Safety Compliance Documentation", "Risk assessment, regulatory compliance scoring"),
]
for r, (num, doc_name, usage) in enumerate(ops_data):
    table.rows[r + 1].cells[0].text = num
    table.rows[r + 1].cells[1].text = doc_name
    table.rows[r + 1].cells[2].text = usage

doc.add_paragraph("")
doc.add_heading("2.3 Strategic Documents (Goals, OKRs & Growth)", level=2)
table = doc.add_table(rows=7, cols=3)
table.style = "Light Grid Accent 1"
headers = ["#", "Document", "How It Is Used"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
strat_data = [
    ("1", "Goals, OKRs & Quarterly Plans", "Goal completion rate, strategy generation, AI recommendations"),
    ("2", "Market Analysis Reports", "Market trend identification, competitive intelligence, growth impact assessment"),
    ("3", "Industry Benchmarks", "Relative performance positioning, gap analysis, target setting"),
    ("4", "Product Roadmaps", "Feature prioritization, resource allocation, timeline planning"),
    ("5", "Technical Architecture Docs", "Tech debt analysis, scalability assessment, modernization roadmap"),
    ("6", "Marketing & Sales Playbook", "Channel effectiveness, CAC analysis, conversion optimization"),
]
for r, (num, doc_name, usage) in enumerate(strat_data):
    table.rows[r + 1].cells[0].text = num
    table.rows[r + 1].cells[1].text = doc_name
    table.rows[r + 1].cells[2].text = usage

doc.add_paragraph("")
doc.add_heading("2.4 HR & Organizational Documents", level=2)
table = doc.add_table(rows=6, cols=3)
table.style = "Light Grid Accent 1"
headers = ["#", "Document", "How It Is Used"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
hr_data = [
    ("1", "Organization Chart & Team List", "Org structure analysis, department breakdown, span of control"),
    ("2", "Employee Roles & Department Definitions", "Role registry, workload distribution, skill gap analysis"),
    ("3", "Staff Scheduling & Utilization Data", "Overload detection, capacity planning, work pattern analysis"),
    ("4", "Check-In Records", "Owner sentiment, engagement tracking, AI persona refinement"),
    ("5", "Journal & Idea Entries", "Innovation pipeline, unworked ideas, mood/wellness trends"),
]
for r, (num, doc_name, usage) in enumerate(hr_data):
    table.rows[r + 1].cells[0].text = num
    table.rows[r + 1].cells[1].text = doc_name
    table.rows[r + 1].cells[2].text = usage

doc.add_paragraph("")
doc.add_heading("3. Generated Reports (Outputs That Define Org Health)", level=1)
doc.add_paragraph(
    "These are the reports and scores YesBoss generates after processing the above documents. "
    "They provide a complete picture of organizational health and growth opportunities."
)

doc.add_heading("3.1 Organizational Health Report", level=2)
health_items = [
    "Health Score (0-100) with label: Healthy / Needs Attention / At Risk",
    "Goal completion rate and task completion rate",
    "Quality score (average quality across all task outcomes)",
    "Overdue/escalated task counts and open bottlenecks",
    "Team size and org structure assessment",
    "Department breakdowns with performance metrics",
    "Work pattern analysis — overloaded employees, best performers per category",
    "Unworked ideas from journal entries with potential impact ranking",
    "Market alignment score",
    "AI-generated strategic recommendations",
]
for item in health_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3.2 Employee Performance Report", level=2)
emp_items = [
    "Total / completed / pending / in-progress / overdue tasks",
    "Escalation status and history",
    "Completion rate and average completion time",
    "Goals contributed to with success rate",
    "Work patterns analysis by category with org comparison",
    "AI-generated feedback and development suggestions",
    "Manager notification status",
]
for item in emp_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3.3 Weekly Performance Report", level=2)
weekly_items = [
    "Executive Summary — active goals, completed goals, task stats, team size, completion rate",
    "Goals Overview table with progress tracking",
    "Recent Tasks table with status and assignee",
    "Department Breakdown table with key metrics",
    "Task Breakdown by Assignee with days open",
    "Employee Insights — name, completed, pending, overdue, stuck, AI suggestion",
    "Recommended Actions based on overdue/stuck employees",
]
for item in weekly_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("3.4 Industry-Specific Dashboard KPIs", level=2)
doc.add_paragraph("Each industry gets tailored metrics across 5 modules:")

table = doc.add_table(rows=6, cols=3)
table.style = "Light Grid Accent 1"
headers = ["Industry", "Sample KPIs", "Modules"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
kpi_data = [
    ("Technology", "Run rate, Burn rate, Runway months, ARR growth, Deployment frequency, Uptime, Sprint velocity, Tech debt", "Founder, Finance, Ops, Productivity, Workflow"),
    ("Finance", "AUM, Client count, Total assets, Net income, ROE, Operational efficiency", "Founder, Finance, Ops, Productivity, Workflow"),
    ("Healthcare", "Patient count, Readmission rate, Bed occupancy, Wait times, Staff utilization", "Founder, Finance, Ops, Productivity, Workflow"),
    ("Retail", "Foot traffic, Conversion rate, LTV, Inventory turnover, Shrink rate", "Founder, Finance, Ops, Productivity, Workflow"),
    ("Manufacturing", "OEE, Downtime, Defect rate, Lead time, Throughput", "Founder, Finance, Ops, Productivity, Workflow"),
]
for r, (industry, kpis, modules) in enumerate(kpi_data):
    table.rows[r + 1].cells[0].text = industry
    table.rows[r + 1].cells[1].text = kpis
    table.rows[r + 1].cells[2].text = modules

doc.add_paragraph("")
doc.add_heading("4. AI-Processed Documents (Continuous Learning Inputs)", level=1)
doc.add_paragraph(
    "These are not required upfront but are continuously ingested by YesBoss to improve "
    "its learning engine and provide increasingly accurate recommendations."
)

ai_items = [
    ("Uploaded Documents", "PDF, Excel, Word, CSV, Images — processed into vector embeddings for semantic search"),
    ("Meeting Notes", "Audio/text meeting uploads — AI extracts tasks, action items, deadlines, participants"),
    ("Journal / Idea Entries", "Voice-to-text or typed entries — mood tracking, goal/task extraction, innovation pipeline"),
    ("Market Trend Data", "AI-fetched industry trends — growth impact assessment, competitive intelligence"),
    ("Workflow Patterns", "Recorded from daily task execution — continuous learning, bottleneck detection"),
    ("Check-In Responses", "Periodic owner surveys — sentiment analysis, strategic alignment check"),
]
table = doc.add_table(rows=len(ai_items) + 1, cols=2)
table.style = "Light Grid Accent 1"
headers = ["Input Type", "Processing & Value"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, (input_type, value) in enumerate(ai_items):
    table.rows[r + 1].cells[0].text = input_type
    table.rows[r + 1].cells[1].text = value

doc.add_paragraph("")
doc.add_heading("5. Summary: Complete Document Matrix", level=1)

summary_table = doc.add_table(rows=9, cols=3)
summary_table.style = "Light Grid Accent 1"
headers = ["Phase / Purpose", "Key Documents", "Business Value"]
for i, h in enumerate(headers):
    summary_table.rows[0].cells[i].text = h
summary_data = [
    ("Onboard Company & Owner", "About us, Org chart, Product list, Pricing sheet, SOPs, OKRs, Contracts, Industry-specific docs", "AI understands the business, sets up correct industry dashboard, generates default goals"),
    ("Financial Health Assessment", "P&L, Balance Sheet, Cash Flow, Tax/Audit reports, Revenue data", "Profitability analysis, burn rate, runway, financial stability score"),
    ("Operational Efficiency", "SOPs, Workflow docs, Quality reports, Maintenance logs, Safety docs", "Bottleneck detection, efficiency scoring, automation opportunities"),
    ("Strategic Growth", "OKRs, Market analysis, Product roadmap, Competitor intel, Sales playbook", "Goal completion tracking, market alignment, growth recommendations"),
    ("Org Health Dashboard", "All of the above + Check-in records, Journal entries, Meeting notes", "Real-time Health Score (0-100), department performance, trend analysis"),
    ("Employee Performance", "Role definitions, Task data, Work patterns, Org hierarchy", "Per-employee productivity, overload detection, AI coaching suggestions"),
    ("Weekly / Periodic Reports", "Live data from all ingested documents", "Executive summaries, action items, recommended interventions"),
    ("AI Continuous Learning", "All uploaded documents, Workflow patterns, Market data", "Cross-company pattern learning, increasingly accurate insights over time"),
]
for r, (phase, docs, value) in enumerate(summary_data):
    summary_table.rows[r + 1].cells[0].text = phase
    summary_table.rows[r + 1].cells[1].text = docs
    summary_table.rows[r + 1].cells[2].text = value

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Note: ")
run.bold = True
p.add_run(
    "The more documents the owner provides during onboarding, the more accurate and valuable "
    "YesBoss's AI-generated insights, goals, KPIs, and recommendations will be. The system is "
    "designed to work with minimal data but improves significantly with each additional document."
)

output_path = "C:\\VSLLP\\krisha\\3\\YesBoss_Document_Requirements.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
