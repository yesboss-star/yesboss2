import logging
from fastapi import APIRouter, HTTPException, Depends, Query
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
from datetime import datetime, timedelta
from ..core.database import get_database
from ..dependencies.auth import get_current_user, get_current_user_optional
from bson import ObjectId
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import io

router = APIRouter()
logger = logging.getLogger("yesboss.reports")

def get_user_org_id(user) -> Optional[str]:
    if hasattr(user, 'user_metadata') and user.user_metadata:
        return user.user_metadata.get("organization_id")
    return None

class ReportRequest(BaseModel):
    period: str = "weekly"
    sections: Optional[List[str]] = None
    organization_id: Optional[str] = None

def build_report_content(db, org_id: str, request: ReportRequest) -> Dict[str, Any]:
    goals = list(db.goals.find({"organization_id": org_id}).sort("created_at", -1))
    tasks = list(db.tasks.find({"organization_id": org_id}).sort("created_at", -1))
    members = list(db.org_chart_members.find({"organization_id": org_id}))

    total_goals = len(goals)
    active_goals = len([g for g in goals if g.get("status") == "active"])
    completed_goals = len([g for g in goals if g.get("status") == "completed"])

    total_tasks = len(tasks)
    completed_tasks = len([t for t in tasks if t.get("status") == "completed"])
    pending_tasks = len([t for t in tasks if t.get("status") == "pending"])
    in_progress_tasks = len([t for t in tasks if t.get("status") == "in_progress"])

    departments = {}
    for g in goals:
        dept = g.get("department", "general")
        if dept not in departments:
            departments[dept] = {"goals": 0, "tasks": 0}
        departments[dept]["goals"] += 1
    for t in tasks:
        dept = t.get("department", "general")
        if dept not in departments:
            departments[dept] = {"goals": 0, "tasks": 0}
        departments[dept]["tasks"] += 1

    return {
        "generated_at": datetime.utcnow().isoformat(),
        "period": request.period,
        "summary": {
            "total_goals": total_goals,
            "active_goals": active_goals,
            "completed_goals": completed_goals,
            "total_tasks": total_tasks,
            "completed_tasks": completed_tasks,
            "pending_tasks": pending_tasks,
            "in_progress_tasks": in_progress_tasks,
            "team_size": len(members),
            "completion_rate": round((completed_tasks / total_tasks * 100) if total_tasks > 0 else 0, 1),
        },
        "departments": departments,
        "goals": [{"title": g.get("title"), "status": g.get("status"), "priority": g.get("priority"), "department": g.get("department")} for g in goals[:10]],
        "tasks": [{"title": t.get("title"), "status": t.get("status"), "priority": t.get("priority")} for t in tasks[:10]],
    }

def generate_pdf(content: Dict[str, Any], org_name: str = "YesBoss") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )
    styles = getSampleStyleSheet()

    primary = colors.HexColor('#0ea5e9')
    dark = colors.HexColor('#1e293b')
    muted = colors.HexColor('#64748b')
    border = colors.HexColor('#e2e8f0')
    light_bg = colors.HexColor('#f8fafc')

    title_style = ParagraphStyle('Title2', parent=styles['Title'], fontSize=26, spaceAfter=4, textColor=dark, leading=32)
    subtitle_style = ParagraphStyle('Sub', parent=styles['Normal'], fontSize=10, spaceAfter=20, textColor=muted, leading=14)
    section_style = ParagraphStyle('Section', parent=styles['Heading2'], fontSize=15, spaceAfter=8, spaceBefore=18, textColor=primary, leading=20)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9.5, spaceAfter=6, leading=14, textColor=dark)
    cell_style = ParagraphStyle('Cell', parent=styles['Normal'], fontSize=8.5, leading=11, textColor=dark)
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, spaceBefore=20, textColor=muted, alignment=1)

    story = []

    story.append(Paragraph(org_name, title_style))
    story.append(Paragraph(f"Weekly Performance Report &mdash; {content.get('period', 'weekly').capitalize()}", subtitle_style))
    story.append(Paragraph(f"Generated: {content['generated_at'][:10]}", subtitle_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph("<hr/>", body_style))
    story.append(Spacer(1, 6))

    summary = content["summary"]

    story.append(Paragraph("Executive Summary", section_style))
    metrics_table = [
        ["Metric", "Value"],
        ["Active Goals", str(summary["active_goals"])],
        ["Completed Goals", str(summary["completed_goals"])],
        ["Total Tasks", str(summary["total_tasks"])],
        ["Completed Tasks", str(summary["completed_tasks"])],
        ["Pending Tasks", str(summary["pending_tasks"])],
        ["In Progress", str(summary["in_progress_tasks"])],
        ["Team Size", str(summary["team_size"])],
        ["Completion Rate", f"{summary['completion_rate']}%"],
    ]
    tbl = Table(metrics_table, colWidths=[3.2*inch, 2.2*inch])
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), primary),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica'),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, border),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light_bg]),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 6))

    if summary["completion_rate"] >= 50:
        status_text = f"Overall completion rate is <b>{summary['completion_rate']}%</b> — your team is on track."
    else:
        status_text = f"Overall completion rate is <b>{summary['completion_rate']}%</b> — review priorities to improve."
    story.append(Paragraph(status_text, body_style))
    story.append(Spacer(1, 12))

    goals = content.get("goals", [])
    if goals:
        story.append(Paragraph("Goals Overview", section_style))
        goal_header = [Paragraph("<b>Title</b>", cell_style), Paragraph("<b>Status</b>", cell_style), Paragraph("<b>Priority</b>", cell_style), Paragraph("<b>Department</b>", cell_style)]
        goal_rows = [goal_header]
        for g in goals:
            goal_rows.append([
                Paragraph(g.get("title", "-")[:55], cell_style),
                Paragraph(g.get("status", "-"), cell_style),
                Paragraph(g.get("priority", "-"), cell_style),
                Paragraph(g.get("department", "-"), cell_style),
            ])
        tbl = Table(goal_rows, colWidths=[3*inch, 0.9*inch, 0.8*inch, 1*inch])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, border),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light_bg]),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 12))

    tasks = content.get("tasks", [])
    if tasks:
        story.append(Paragraph("Recent Tasks", section_style))
        task_header = [Paragraph("<b>Title</b>", cell_style), Paragraph("<b>Status</b>", cell_style), Paragraph("<b>Priority</b>", cell_style)]
        task_rows = [task_header]
        for t in tasks:
            task_rows.append([
                Paragraph(t.get("title", "-")[:55], cell_style),
                Paragraph(t.get("status", "-"), cell_style),
                Paragraph(t.get("priority", "-"), cell_style),
            ])
        tbl = Table(task_rows, colWidths=[3.8*inch, 1*inch, 1*inch])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, border),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light_bg]),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 12))

    departments = content.get("departments", {})
    if departments:
        story.append(Paragraph("Department Breakdown", section_style))
        dept_header = [Paragraph("<b>Department</b>", cell_style), Paragraph("<b>Goals</b>", cell_style), Paragraph("<b>Tasks</b>", cell_style)]
        dept_rows = [dept_header]
        for dept, info in departments.items():
            dept_rows.append([
                Paragraph(dept.capitalize(), cell_style),
                Paragraph(str(info.get("goals", 0)), cell_style),
                Paragraph(str(info.get("tasks", 0)), cell_style),
            ])
        tbl = Table(dept_rows, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 7),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 0.5, border),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light_bg]),
        ]))
        story.append(tbl)

    story.append(Spacer(1, 30))
    story.append(Paragraph("<hr/>", body_style))
    story.append(Paragraph("This report was generated automatically by YesBoss AI Business Operating System.", footer_style))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()

@router.post("/generate")
async def generate_report(
    request: ReportRequest,
    current_user = Depends(get_current_user_optional)
):
    db = get_database()
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")

    org_id = request.organization_id or get_user_org_id(current_user)
    if not org_id:
        raise HTTPException(status_code=400, detail="Organization ID required")

    goals_count = db.goals.count_documents({"organization_id": org_id})
    tasks_count = db.tasks.count_documents({"organization_id": org_id})

    if goals_count == 0 and tasks_count == 0:
        raise HTTPException(
            status_code=400,
            detail="Insufficient data to generate report. Add goals and tasks first."
        )

    content = build_report_content(db, org_id, request)

    report_doc = {
        "organization_id": org_id,
        "period": request.period,
        "content": content,
        "created_at": datetime.utcnow(),
        "created_by": getattr(current_user, 'id', None),
    }

    result = db.reports.insert_one(report_doc)
    report_doc["_id"] = str(result.inserted_id)

    return {
        "report": {
            "id": report_doc["_id"],
            "period": request.period,
            "generated_at": content["generated_at"],
            "summary": content["summary"],
            "departments": content["departments"],
        }
    }

@router.get("/history")
async def list_reports(current_user = Depends(get_current_user)):
    db = get_database()
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")

    org_id = get_user_org_id(current_user)
    if not org_id:
        raise HTTPException(status_code=400, detail="Organization ID required")

    reports = list(db.reports.find({"organization_id": org_id}).sort("created_at", -1).limit(20))
    for r in reports:
        r["_id"] = str(r["_id"])
        if "content" in r and isinstance(r["content"], dict):
            r["summary"] = r["content"].get("summary", {})
    return {"reports": reports}

def generate_docx(content: Dict[str, Any], org_name: str = "YesBoss") -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    title = doc.add_heading(org_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    period_str = content.get('period', 'weekly').capitalize()
    doc.add_paragraph(f"Weekly Performance Report — {period_str}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {content['generated_at'][:10]}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_heading("Executive Summary", level=1)
    summary = content["summary"]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Value'

    metrics_rows = [
        ("Active Goals", str(summary["active_goals"])),
        ("Completed Goals", str(summary["completed_goals"])),
        ("Total Tasks", str(summary["total_tasks"])),
        ("Completed Tasks", str(summary["completed_tasks"])),
        ("Pending Tasks", str(summary["pending_tasks"])),
        ("In Progress", str(summary["in_progress_tasks"])),
        ("Team Size", str(summary["team_size"])),
        ("Completion Rate", f"{summary['completion_rate']}%"),
    ]
    for metric, value in metrics_rows:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = value

    doc.add_paragraph("")
    if summary['completion_rate'] >= 50:
        status_text = f"Overall completion rate is {summary['completion_rate']}% — your team is on track."
    else:
        status_text = f"Overall completion rate is {summary['completion_rate']}% — review priorities to improve."
    doc.add_paragraph(status_text)

    goals = content.get("goals", [])
    if goals:
        doc.add_heading("Goals Overview", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Shading Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Title'
        hdr[1].text = 'Status'
        hdr[2].text = 'Priority'
        hdr[3].text = 'Department'
        for g in goals:
            row = table.add_row().cells
            row[0].text = g.get("title", "-")[:55]
            row[1].text = g.get("status", "-")
            row[2].text = g.get("priority", "-")
            row[3].text = g.get("department", "-")

    tasks = content.get("tasks", [])
    if tasks:
        doc.add_heading("Recent Tasks", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Title'
        hdr[1].text = 'Status'
        hdr[2].text = 'Priority'
        for t in tasks:
            row = table.add_row().cells
            row[0].text = t.get("title", "-")[:55]
            row[1].text = t.get("status", "-")
            row[2].text = t.get("priority", "-")

    departments = content.get("departments", {})
    if departments:
        doc.add_heading("Department Breakdown", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Department'
        hdr[1].text = 'Goals'
        hdr[2].text = 'Tasks'
        for dept, info in departments.items():
            row = table.add_row().cells
            row[0].text = dept.capitalize()
            row[1].text = str(info.get("goals", 0))
            row[2].text = str(info.get("tasks", 0))

    doc.add_paragraph("")
    doc.add_paragraph("This report was generated automatically by YesBoss AI Business Operating System.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@router.get("/download/{report_id}")
async def download_report(
    report_id: str,
    format: str = Query("pdf", regex="^(pdf|docx)$"),
    current_user = Depends(get_current_user_optional)
):
    db = get_database()
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")

    report = db.reports.find_one({"_id": ObjectId(report_id)})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    content = report.get("content", {})
    org_id = report.get("organization_id")
    org = db.organizations.find_one({"_id": org_id}) if org_id else None
    org_name = org.get("name", "YesBoss") if org else "YesBoss"

    from fastapi.responses import Response

    if format == "docx":
        docx_bytes = generate_docx(content, org_name)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=yesboss_report_{report_id}.docx",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            }
        )

    pdf_bytes = generate_pdf(content, org_name)
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=yesboss_report_{report_id}.pdf",
            "Content-Type": "application/pdf",
        }
    )
