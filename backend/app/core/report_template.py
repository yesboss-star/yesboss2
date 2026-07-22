import io
import logging
from typing import Any

from docx import Document

logger = logging.getLogger("yesboss.report_template")


PLACEHOLDERS = [
    {"key": "org_name", "description": "Organization name"},
    {"key": "report_date", "description": "Report generation date (YYYY-MM-DD)"},
    {"key": "period", "description": "Report period (weekly, monthly, etc.)"},
    {"key": "active_goals", "description": "Number of active goals"},
    {"key": "completed_goals", "description": "Number of completed goals"},
    {"key": "total_tasks", "description": "Total number of tasks"},
    {"key": "completed_tasks", "description": "Number of completed tasks"},
    {"key": "pending_tasks", "description": "Number of pending tasks"},
    {"key": "in_progress_tasks", "description": "Number of in-progress tasks"},
    {"key": "team_size", "description": "Total team members"},
    {"key": "completion_rate", "description": "Task completion rate (percentage)"},
    {"key": "executive_summary", "description": "AI-written executive summary paragraph"},
]


def list_placeholders() -> list[dict]:
    return PLACEHOLDERS


def fill_template(template_bytes: bytes, data: dict) -> bytes:
    doc = Document(io.BytesIO(template_bytes))

    summary = data.get("summary", {})
    content = data.get("content", data)

    replacements = {
        "{{org_name}}": data.get("org_name", ""),
        "{{report_date}}": content.get("generated_at", "")[:10] if content.get("generated_at") else "",
        "{{period}}": content.get("period", "weekly").capitalize(),
        "{{active_goals}}": str(summary.get("active_goals", 0)),
        "{{completed_goals}}": str(summary.get("completed_goals", 0)),
        "{{total_tasks}}": str(summary.get("total_tasks", 0)),
        "{{completed_tasks}}": str(summary.get("completed_tasks", 0)),
        "{{pending_tasks}}": str(summary.get("pending_tasks", 0)),
        "{{in_progress_tasks}}": str(summary.get("in_progress_tasks", 0)),
        "{{team_size}}": str(summary.get("team_size", 0)),
        "{{completion_rate}}": f"{summary.get('completion_rate', 0)}%",
        "{{executive_summary}}": _build_executive_summary(summary),
    }

    for para in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in para.text:
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                para.text = para.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
                            para.text = para.text.replace(placeholder, str(value))

    _fill_dynamic_tables(doc, content)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_executive_summary(summary: dict) -> str:
    rate = summary.get("completion_rate", 0)
    active = summary.get("active_goals", 0)
    completed = summary.get("completed_tasks", 0)
    total = summary.get("total_tasks", 0)
    team = summary.get("team_size", 0)

    parts = [
        f"This week, your team completed {completed} of {total} tasks "
        f"({rate}% completion rate)"
    ]
    if active:
        parts.append(f"with {active} active goal(s) in progress")
    if team:
        parts.append(f"across {team} team members")
    parts.append(".")
    return " ".join(parts)


def _fill_dynamic_tables(doc: Document, content: dict):
    goals = content.get("goals", [])
    tasks = content.get("tasks", [])
    departments = content.get("departments", {})

    for table in doc.tables:
        first_cell_text = ""
        for row in table.rows:
            for cell in row.cells:
                first_cell_text = cell.text.strip()
                break
            if first_cell_text:
                break

        if "{{goals_table}}" in table._tbl.xml if hasattr(table, '_tbl') else False:
            pass

        if first_cell_text == "{{goals_table}}" or "{{goals_table}}" in table._tbl.xml if hasattr(table, '_tbl') else False:
            _replace_goals_table(table, goals)
        elif first_cell_text == "{{tasks_table}}" or "{{tasks_table}}" in table._tbl.xml if hasattr(table, '_tbl') else False:
            _replace_tasks_table(table, tasks)
        elif first_cell_text == "{{department_breakdown}}" or "{{department_breakdown}}" in table._tbl.xml if hasattr(table, '_tbl') else False:
            _replace_department_table(table, departments)
        elif first_cell_text == "{{employee_insights}}" or "{{employee_insights}}" in table._tbl.xml if hasattr(table, '_tbl') else False:
            _replace_employee_table(table, content.get("employee_insights", []))
        elif first_cell_text == "{{task_breakdown}}" or "{{task_breakdown}}" in table._tbl.xml if hasattr(table, '_tbl') else False:
            _replace_task_breakdown_table(table, content.get("task_breakdown", []))


def _replace_goals_table(table, goals: list[dict]):
    _clear_table_rows(table)
    headers = ["Title", "Status", "Priority", "Department"]
    _add_header_row(table, headers)
    for g in goals:
        row = table.add_row().cells
        row[0].text = g.get("title", "-")[:55]
        row[1].text = g.get("status", "-")
        row[2].text = g.get("priority", "-")
        row[3].text = g.get("department", "-")


def _replace_tasks_table(table, tasks: list[dict]):
    _clear_table_rows(table)
    headers = ["Title", "Status", "Priority"]
    _add_header_row(table, headers)
    for t in tasks:
        row = table.add_row().cells
        row[0].text = t.get("title", "-")[:55]
        row[1].text = t.get("status", "-")
        row[2].text = t.get("priority", "-")


def _replace_department_table(table, departments: dict):
    _clear_table_rows(table)
    headers = ["Department", "Goals", "Tasks"]
    _add_header_row(table, headers)
    for dept, info in departments.items():
        row = table.add_row().cells
        row[0].text = dept.capitalize()
        row[1].text = str(info.get("goals", 0))
        row[2].text = str(info.get("tasks", 0))


def _replace_employee_table(table, insights: list[dict]):
    _clear_table_rows(table)
    headers = ["Employee", "Done", "Pending", "Overdue", "Stuck", "Suggestion"]
    _add_header_row(table, headers)
    for e in insights:
        row = table.add_row().cells
        row[0].text = e.get("name", "-")[:20]
        row[1].text = str(e.get("completed", 0))
        row[2].text = str(e.get("pending", 0))
        row[3].text = str(e.get("overdue", 0))
        row[4].text = str(e.get("stuck", 0))
        row[5].text = e.get("suggestion", "")[:40]


def _replace_task_breakdown_table(table, breakdown: list[dict]):
    _clear_table_rows(table)
    headers = ["Task", "Status", "Assignee", "Priority", "Days"]
    _add_header_row(table, headers)
    for t in breakdown:
        row = table.add_row().cells
        row[0].text = t.get("title", "-")
        row[1].text = t.get("status", "-")
        row[2].text = t.get("assignee", "-")[:20]
        row[3].text = t.get("priority", "-")
        row[4].text = str(t.get("days_open", 0))


def _clear_table_rows(table):
    for _ in range(len(table.rows)):
        table._tbl.remove(table._tbl.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')[0])  # type: ignore[attr-defined]


def _add_header_row(table, headers: list[str]):
    row = table.add_row().cells
    for i, h in enumerate(headers):
        if i < len(row):
            row[i].text = h
