import io
import re
import logging
import uuid
from typing import Optional, List, Dict, Any
from datetime import datetime

logger = logging.getLogger("yesboss.file_processor")

ALLOWED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".png", ".jpg", ".jpeg", ".docx", ".txt", ".csv"}

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


def get_file_type(filename: str) -> str:
    ext = filename.lower().split(".")[-1] if "." in filename else ""
    type_map = {
        "pdf": "pdf",
        "xlsx": "excel",
        "xls": "excel",
        "png": "image",
        "jpg": "image",
        "jpeg": "image",
        "docx": "word",
        "txt": "text",
        "csv": "csv",
    }
    return type_map.get(ext, "unknown")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, doc_type="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text.strip()
        except Exception as e2:
            logger.error(f"PDF fallback also failed: {e2}")
            return ""


def extract_text_from_excel(file_bytes: bytes) -> str:
    try:
        import pandas as pd
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        text = ""
        for sheet_name, sheet_df in df.items():
            text += f"\n## Sheet: {sheet_name}\n"
            text += sheet_df.to_string() + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Excel extraction failed: {e}")
        return ""


def extract_text_from_image(file_bytes: bytes, file_type: str) -> str:
    try:
        from PIL import Image
        import pytesseract
        
        image = Image.open(io.BytesIO(file_bytes))
        if image.mode != "RGB":
            image = image.convert("RGB")
        
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        try:
            from transformers import pipeline
            import torch
            
            if torch.cuda.is_available():
                ocr_pipeline = pipeline("ocr", model="microsoft/trocr-base-handwritten", device=0)
            else:
                ocr_pipeline = pipeline("ocr", model="microsoft/trocr-base-handwritten")
            
            image = Image.open(io.BytesIO(file_bytes))
            result = ocr_pipeline(image)
            return result[0]["generated_text"] if result else ""
        except Exception as e2:
            logger.error(f"OCR fallback also failed: {e2}")
            return ""


def extract_text_from_word(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Word extraction failed: {e}")
        return ""


def extract_text_from_csv(file_bytes: bytes) -> str:
    try:
        import pandas as pd
        df = pd.read_csv(io.BytesIO(file_bytes))
        return df.to_string()
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        return ""


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return ""


def extract_text(file_bytes: bytes, filename: str) -> str:
    file_type = get_file_type(filename)
    logger.info(f"Extracting text from {filename} (type: {file_type})")
    
    extractors = {
        "pdf": extract_text_from_pdf,
        "excel": extract_text_from_excel,
        "image": lambda b: extract_text_from_image(b, "image"),
        "word": extract_text_from_word,
        "csv": extract_text_from_csv,
        "text": extract_text_from_txt,
    }
    
    extractor = extractors.get(file_type)
    if extractor:
        return extractor(file_bytes)
    
    return ""


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    if not text:
        return []
    
    chunks = []
    for i in range(0, len(text), chunk_size - overlap):
        chunk = text[i:i + chunk_size]
        if chunk.strip():
            chunks.append(chunk)
    return chunks


async def generate_embeddings(texts: List[str], provider: Optional[str] = None) -> List[List[float]]:
    from .config import settings
    resolved_provider = provider or settings.DEFAULT_AI_PROVIDER
    
    embeddings = []
    
    try:
        from openai import AsyncOpenAI
        
        if resolved_provider == "xai":
            api_key = settings.XAI_API_KEY
            base_url = settings.XAI_BASE_URL
        elif resolved_provider == "openai":
            api_key = settings.OPENAI_API_KEY
            base_url = None
        elif resolved_provider == "qwen":
            api_key = settings.QWEN_API_KEY
            base_url = settings.QWEN_BASE_URL
        else:
            api_key = None
            base_url = None
        
        if api_key:
            client = AsyncOpenAI(api_key=api_key, base_url=base_url)
            for text in texts:
                response = await client.embeddings.create(
                    model="text-embedding-3-small",
                    input=text[:8000]
                )
                embeddings.append(response.data[0].embedding)
        else:
            for text in texts:
                embeddings.append([0.0] * 1536)
    except Exception as e:
        logger.error(f"Embedding failed with {resolved_provider}: {e}")
        for text in texts:
            embeddings.append([0.0] * 1536)
    
    return embeddings


async def store_embeddings_in_qdrant(
    collection_name: str,
    vectors: List[List[float]],
    payloads: List[Dict]
):
    from .qdrant import get_qdrant_client
    
    try:
        client = get_qdrant_client()
        
        points = []
        for i, (vector, payload) in enumerate(zip(vectors, payloads)):
            points.append({
                "id": str(uuid.uuid4()),
                "vector": vector,
                "payload": payload
            })
        
        if points:
            client.upsert(
                collection_name=collection_name,
                points=points
            )
        
        return True
    except Exception as e:
        logger.error(f"Failed to store in Qdrant: {e}")
        return False


async def process_file(
    file_bytes: bytes,
    filename: str,
    org_id: str,
    user_id: str,
    provider: Optional[str] = None,
    company_name: str = "",
    industry: str = "",
    micro_vertical: str = "",
) -> Dict[str, Any]:
    logger.info(f"Processing file: {filename} for org: {org_id}")

    file_id = str(uuid.uuid4())
    file_type = get_file_type(filename)

    text = extract_text(file_bytes, filename)

    if not text:
        return {
            "file_id": file_id,
            "filename": filename,
            "status": "failed",
            "error": "Could not extract text from file",
        }

    chunks = chunk_text(text)

    metadata = {
        "file_id": file_id,
        "filename": filename,
        "file_type": file_type,
        "org_id": org_id,
        "user_id": user_id,
        "text_length": len(text),
        "chunk_count": len(chunks),
        "processed_at": datetime.utcnow().isoformat(),
        "document_type": file_type,
    }

    from .database import get_database
    db = get_database()

    doc = {
        "file_id": file_id,
        "filename": filename,
        "file_type": file_type,
        "org_id": org_id,
        "user_id": user_id,
        "text": text[:10000],
        "text_length": len(text),
        "chunks": chunks,
        "chunk_count": len(chunks),
        "insights": None,
        "insights_status": "pending",
        "created_at": datetime.utcnow(),
        "metadata": metadata,
    }

    try:
        db.documents.insert_one(doc)
    except Exception as e:
        logger.error(f"Failed to save document to MongoDB: {e}")

    if chunks:
        try:
            embeddings = await generate_embeddings(chunks, provider)

            payloads = []
            for i, chunk in enumerate(chunks):
                payloads.append({
                    "file_id": file_id,
                    "filename": filename,
                    "org_id": org_id,
                    "chunk_index": i,
                    "chunk_text": chunk[:500],
                    "document_type": file_type,
                    "processed_at": datetime.utcnow().isoformat(),
                })

            await store_embeddings_in_qdrant("documents", embeddings, payloads)
        except Exception as e:
            logger.error(f"Failed to create embeddings: {e}")

    import asyncio
    asyncio.create_task(
        _run_deep_analysis(
            file_id=file_id,
            org_id=org_id,
            filename=filename,
            file_type=file_type,
            text=text,
            company_name=company_name,
            industry=industry,
            micro_vertical=micro_vertical,
        )
    )

    return {
        "file_id": file_id,
        "filename": filename,
        "file_type": file_type,
        "status": "completed",
        "text_length": len(text),
        "chunks": len(chunks),
        "insights_status": "pending",
        "message": "File processed successfully. Deep analysis running in background.",
    }


async def _run_deep_analysis(
    file_id: str,
    org_id: str,
    filename: str,
    file_type: str,
    text: str,
    company_name: str,
    industry: str,
    micro_vertical: str,
) -> None:
    """Background task: extract structured insights and write back to MongoDB."""
    try:
        from .intelligence import extract_document_insights
        from .database import get_database

        insights = await extract_document_insights(
            text=text,
            filename=filename,
            file_type=file_type,
            company_name=company_name,
            industry=industry,
            micro_vertical=micro_vertical,
        )

        db = get_database()
        if db is not None:
            db.documents.update_one(
                {"file_id": file_id, "org_id": org_id},
                {
                    "$set": {
                        "insights": insights,
                        "insights_status": "completed",
                        "insights_at": datetime.utcnow(),
                    }
                },
            )
        logger.info(f"Deep analysis complete for {filename} ({file_id})")
    except Exception as e:
        logger.error(f"Deep analysis failed for {file_id}: {e}")
        try:
            from .database import get_database
            db = get_database()
            if db is not None:
                db.documents.update_one(
                    {"file_id": file_id, "org_id": org_id},
                    {
                        "$set": {
                            "insights_status": "failed",
                            "insights_error": str(e),
                            "insights_at": datetime.utcnow(),
                        }
                    },
                )
        except Exception:
            pass


async def get_org_document_context(
    org_id: str,
    max_docs: int = 20,
) -> Dict[str, Any]:
    """Return the structured context (summaries, metrics, decisions) of all
    analyzed documents for an organization. Powers dashboard widgets and the
    AI assistant's system prompt.
    """
    from .database import get_database

    db = get_database()
    if db is None:
        return {"documents": [], "summary": "", "metrics": []}

    try:
        docs = list(
            db.documents.find({"org_id": org_id})
            .sort("created_at", -1)
            .limit(max_docs)
        )
    except Exception as e:
        logger.error(f"Failed to load documents for org {org_id}: {e}")
        return {"documents": [], "summary": "", "metrics": []}

    out_docs = []
    aggregated_metrics = []
    category_counts: Dict[str, int] = {}
    summaries_for_brief: list = []

    for d in docs:
        insights = d.get("insights") or {}
        cat = insights.get("document_category") or d.get("file_type") or "other"
        category_counts[cat] = category_counts.get(cat, 0) + 1

        for m in insights.get("key_metrics", []) or []:
            if isinstance(m, dict) and m.get("name"):
                aggregated_metrics.append({
                    "source_file": d.get("filename"),
                    "name": m.get("name", ""),
                    "value": m.get("value", ""),
                    "context": m.get("context", ""),
                })

        if insights.get("summary"):
            summaries_for_brief.append(f"- {d.get('filename')}: {insights['summary']}")

        out_docs.append({
            "file_id": d.get("file_id"),
            "filename": d.get("filename"),
            "file_type": d.get("file_type"),
            "uploaded_at": d.get("created_at").isoformat() if d.get("created_at") else None,
            "insights_status": d.get("insights_status", "pending"),
            "summary": insights.get("summary", ""),
            "document_category": cat,
            "key_metrics": insights.get("key_metrics", []),
            "key_entities": insights.get("key_entities", {}),
            "decisions": insights.get("decisions", []),
            "action_items": insights.get("action_items", []),
            "qa_pairs": insights.get("qa_pairs", []),
        })

    brief_lines = []
    if out_docs:
        brief_lines.append(f"Total documents: {len(out_docs)}")
        if category_counts:
            cats = ", ".join(f"{k}: {v}" for k, v in sorted(category_counts.items(), key=lambda x: -x[1])[:5])
            brief_lines.append(f"Categories: {cats}")
        brief_lines.append("")
        brief_lines.extend(summaries_for_brief[:10])

    return {
        "documents": out_docs,
        "summary": "\n".join(brief_lines).strip(),
        "metrics": aggregated_metrics[:50],
        "category_breakdown": category_counts,
        "total_documents": len(out_docs),
        "analyzed_documents": sum(1 for d in out_docs if d["insights_status"] == "completed"),
        "pending_documents": sum(1 for d in out_docs if d["insights_status"] == "pending"),
    }


async def search_documents(
    org_id: str,
    query: str,
    top_k: int = 5,
    provider: Optional[str] = None
) -> List[Dict[str, Any]]:
    logger.info(f"Searching documents for org: {org_id}, query: {query}")

    from .ai_client import get_ai_response

    query_vector = None
    try:
        query_embedding = await generate_embeddings([query], provider)
        query_vector = query_embedding[0] if query_embedding and not all(v == 0.0 for v in query_embedding[0]) else None
    except Exception as e:
        logger.warning(f"Embedding generation failed for search: {e}")
        query_vector = None
    
    if query_vector is not None:
        try:
            from .qdrant import get_qdrant_client
            from qdrant_client.models import Filter, FieldCondition, Match, MatchExcept

            client = get_qdrant_client()

            results = client.search(
                collection_name="documents",
                query_vector=query_vector,
                limit=top_k,
                query_filter=Filter(
                    must=[FieldCondition(key="org_id", match=Match(value=org_id))]
                ),
                with_payload=True
            )

            search_results = []
            for result in results:
                search_results.append({
                    "filename": result.payload.get("filename"),
                    "text": result.payload.get("chunk_text", ""),
                    "score": result.score,
                    "file_id": result.payload.get("file_id")
                })

            if search_results:
                return search_results
        except Exception as e:
            logger.warning(f"Qdrant search failed, falling back to Mongo: {e}")

    from .database import get_database
    db = get_database()
    if db is None:
        return []
    try:
        keywords = [w for w in re.findall(r"[A-Za-z0-9$%]{3,}", query) if w.lower() not in {"the","and","for","with","what","our","how","did","was","are","this","that","from"}][:8]
        if keywords:
            regex = "|".join(re.escape(k) for k in keywords)
            cursor = db.documents.find(
                {"org_id": org_id, "text": {"$regex": regex, "$options": "i"}},
                {"filename": 1, "text": 1, "file_id": 1, "text_length": 1},
            ).limit(top_k * 3)
        else:
            cursor = db.documents.find(
                {"org_id": org_id},
                {"filename": 1, "text": 1, "file_id": 1, "text_length": 1},
            ).limit(top_k)

        out = []
        for d in cursor:
            txt = (d.get("text") or "")[:1200]
            score = 0.5
            if keywords:
                lowered = txt.lower()
                hits = sum(1 for k in keywords if k.lower() in lowered)
                score = min(0.99, 0.4 + 0.1 * hits)
            out.append({
                "filename": d.get("filename"),
                "text": txt,
                "score": score,
                "file_id": d.get("file_id"),
            })
        out.sort(key=lambda r: -r["score"])
        return out[:top_k]
    except Exception as e:
        logger.error(f"MongoDB fallback search failed: {e}")
        return []